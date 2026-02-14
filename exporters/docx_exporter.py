from __future__ import annotations

import io

from docx import Document


# Converte título e texto simples em bytes de um arquivo DOCX.
def texto_para_docx_bytes(titulo: str, texto: str) -> bytes:
    doc = Document()
    doc.add_heading(titulo, level=1)

    for linha in texto.splitlines():
        if linha.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(linha)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# Backward-compatible alias used by earlier app versions.
# Mantém compatibilidade com versões antigas que usam o nome em inglês.
def build_docx_bytes(text: str, title: str = "PETICAO INICIAL") -> bytes:
    return texto_para_docx_bytes(titulo=title, texto=text)
