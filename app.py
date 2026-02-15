from __future__ import annotations

import os
import re
import html
import json
import io
import hashlib
from typing import Any
from urllib import error as urlerror
from urllib import request as urlrequest

import streamlit as st
from dotenv import load_dotenv

from exporters.docx_exporter import texto_para_docx_bytes
from exporters.pdf_exporter import texto_para_pdf_bytes
from services.gemini_service import GeminiServiceError, gerar_peticao
from services.prompt_builder import montar_prompt

# ============================================================================
# SISTEMA DE AUTENTICAÇÃO
# ============================================================================

load_dotenv()

SENHA_APP = os.getenv("APP_PASSWORD")  # coloque no Secrets do Streamlit


def exigir_senha():
    """Exige autenticação por senha antes de permitir acesso ao sistema."""
    if "autenticado" not in st.session_state:
        st.session_state.autenticado = False

    if st.session_state.autenticado:
        return True

    st.title("🔒 Acesso restrito")
    st.markdown("---")
    senha = st.text_input("Senha", type="password")

    if st.button("Entrar"):
        if senha == SENHA_APP and SENHA_APP:
            st.session_state.autenticado = True
            st.rerun()
        else:
            st.error("❌ Senha incorreta.")
    return False


# Bloqueia execução se não estiver autenticado
if not exigir_senha():
    st.stop()

# ============================================================================
# CONSTANTES E CONFIGURAÇÕES
# ============================================================================

AREAS_DIREITO = [
    "Previdenciário",
    "Direito da Saúde",
]

ALIAS_AREA_CAMPOS = {
    "Previdenciário": "Previdenciario",
    "Direito da Saúde": "Direito da Saude",
}

TIPOS_ACAO_POR_AREA: dict[str, list[str]] = {
    "Previdenciário": [
        "Concessão de benefício",
        "Restabelecimento de benefício",
        "Revisão de benefício",
        "Auxílio-doença / Benefício por incapacidade",
        "BPC/LOAS",
        "Aposentadoria (idade/tempo/especial)",
        "Mandado de segurança previdenciário",
    ],
    "Direito da Saúde": [
        "Obrigação de fazer (Plano de saúde: cobertura/tratamento/medicamento)",
        "Obrigação de fazer (SUS/Ente público: fornecimento de medicamento/terapia)",
        "Tutela de urgência (tratamento imediato)",
        "Reembolso de despesas médicas",
        "Home care",
        "Internação/UTI",
        "Mandado de segurança (saúde)",
    ],
    "Outro": [
        "Ação ordinária",
        "Ação declaratória",
        "Mandado de segurança",
        "Ação de execução",
    ],
}

RITOS_POR_AREA: dict[str, list[str]] = {
    "Previdenciário": [
        "Juizado Especial Federal (até 60 salários mínimos)",
        "Procedimento Comum (CPC)",
        "Mandado de Segurança (rito próprio)",
    ],
    "Direito da Saúde": [
        "Procedimento Comum (CPC)",
        "Tutela de urgência antecedente (CPC)",
        "Mandado de Segurança (rito próprio)",
        "Juizado Especial (se cabível)",
    ],
    "Outro": [
        "Procedimento Comum (CPC)",
        "Procedimento Especial",
        "Execução",
    ],
}

NATUREZA_RELACAO_OPCOES: dict[str, list[str]] = {
    "Previdenciario": [
        "Benefício por incapacidade",
        "Aposentadoria",
        "BPC/LOAS",
        "Pensão por morte",
        "Outro",
    ],
    "Direito da Saude": [
        "Plano de saúde (contrato/cobertura)",
        "SUS / Ente público (obrigação estatal)",
        "Hospital/Clínica (prestação de serviço)",
        "Profissional de saúde (responsabilidade civil)",
        "Outro",
    ],
    "Outro": [
        "Outro",
    ],
}

PROVAS_SUGERIDAS_POR_AREA: dict[str, list[str]] = {
    "Previdenciário": [
        "CNIS",
        "Carta de indeferimento administrativo",
        "Comprovante de requerimento administrativo",
        "Laudos e relatórios médicos",
        "Atestados médicos",
        "Carteira de trabalho (CTPS)",
        "Comprovantes de contribuição",
    ],
    "Direito da Saúde": [
        "Relatório médico",
        "Prescrição médica",
        "Negativa formal de cobertura/atendimento",
        "Protocolos de atendimento",
        "Notas fiscais e orçamentos",
        "Exames e laudos complementares",
        "Contrato/carteirinha do plano",
    ],
    "Outro": [
        "Documentos essenciais do caso",
    ],
}

TEMAS_JURIDICOS_COMUNS = [
    "Responsabilidade civil",
    "Inadimplemento contratual",
    "Boa-fé objetiva",
    "Dano moral",
    "Dano material",
    "Inversão do ônus da prova",
    "CDC",
    "Tutela de urgência",
    "Justiça gratuita",
    "Prescrição/decadência",
]

PEDIDOS_BASE = [
    "Citação da parte ré",
    "Procedência total",
    "Procedência parcial",
    "Danos morais",
    "Danos materiais",
    "Obrigação de fazer",
    "Obrigação de não fazer",
    "Tutela de urgência",
    "Justiça gratuita",
    "Condenação em custas e honorários",
    "Produção de provas",
]

SECOES_SUGERIDAS = [
    "Dos fatos",
    "Da competência",
    "Do direito",
    "Da tutela de urgência",
    "Dos pedidos",
    "Do valor da causa",
    "Dos requerimentos finais",
]

NIVEIS_DETALHAMENTO = ["Enxuto", "Padrão", "Aprofundado"]
TIPOS_PESSOA_OPCOES = ["Pessoa Física", "Pessoa Jurídica"]
LIMITE_CARACTERES_MODELO_REFERENCIA = 12000
MODO_PREENCHIMENTO_OPCOES = ["Essencial", "Completo"]
PEDIDOS_PARAMETROS_FINAIS = ["Tutela de urgência", "Justiça gratuita"]

ETAPAS_FLUXO = [
    "Contexto Processual",
    "Campos da Área",
    "Partes",
    "Fatos e Provas",
    "Fundamentação",
    "Pedidos",
    "Finalização e Geração",
]

CHAVES_FORMULARIO_BASE = [
    "area_direito",
    "modo_preenchimento",
    "tipo_acao",
    "rito",
    "comarca_uf",
    "foro_vara",
    "autor_tipo_pessoa",
    "autor_nome",
    "autor_doc",
    "autor_cep",
    "autor_end",
    "autor_nacionalidade",
    "autor_estado_civil",
    "autor_profissao",
    "autor_natureza_juridica",
    "autor_representante_legal",
    "autor_qualificacao",
    "reu_tipo_pessoa",
    "reu_nome",
    "reu_doc",
    "reu_cep",
    "reu_end",
    "reu_nacionalidade",
    "reu_estado_civil",
    "reu_profissao",
    "reu_natureza_juridica",
    "reu_representante_legal",
    "reu_qualificacao",
    "partes_adicionais_raw",
    "fatos",
    "cronologia_raw",
    "provas_sugeridas",
    "provas_raw",
    "teses_juridicas",
    "temas_comuns",
    "fundamentos_legais_raw",
    "temas_custom_raw",
    "pedidos_base",
    "pedidos_custom_raw",
    "secoes_sugeridas",
    "secoes_extras_raw",
    "valor_causa",
    "advogado_nome",
    "advogado_oab_uf",
    "advogado_oab_num",
    "nivel_detalhamento",
    "tem_tutela_urgencia",
    "tem_gratuidade",
    "tem_prioridade",
    "quer_audiencia",
    "obs_estrategicas",
    "modelo_referencia_nome",
    "modelo_referencia_texto",
    "modelo_referencia_truncado",
]

CAMPOS_POR_AREA: dict[str, list[dict[str, Any]]] = {
    "Previdenciario": [
        {
            "id": "natureza_relacao_juridica",
            "label": "Natureza da relação jurídica",
            "widget": "select",
            "options": NATUREZA_RELACAO_OPCOES["Previdenciario"],
        },
        {
            "id": "beneficio_pretendido",
            "label": "Benefício pretendido",
            "widget": "select",
            "options": [
                "Aposentadoria por idade",
                "Aposentadoria por tempo de contribuição",
                "Auxílio por incapacidade temporária",
                "BPC/LOAS",
                "Pensão por morte",
                "Aposentadoria por invalidez",
                "Outro",
            ],
        },
        {
            "id": "nb_ou_requerimento",
            "label": "NB/protocolo administrativo (opcional)",
            "widget": "text",
            "placeholder": "Ex.: NB 123.456.789-0",
        },
        {
            "id": "der_dib",
            "label": "DER/DIB (se houver)",
            "widget": "text",
            "placeholder": "Ex.: DER 10/01/2026 - DIB 15/02/2026",
        },
        {
            "id": "tempo_contribuicao_total",
            "label": "Tempo total de contribuição (aproximado)",
            "widget": "select",
            "options": [
                "Até 1 ano",
                "1 a 5 anos",
                "5 a 10 anos",
                "10 a 15 anos",
                "15 anos ou mais",
                "20 anos ou mais",
                "30 anos ou mais",
                "35 anos ou mais",
                "Outro",
            ],
            "accept_new_options": True,
        },
        {
            "id": "qualidade_segurado",
            "label": "Situação da qualidade de segurado",
            "widget": "select",
            "options": [
                "Mantida (contribuindo)",
                "Período de graça – 12 meses",
                "Período de graça – 24 meses",
                "Período de graça – 36 meses (desemprego comprovado)",
                "Em gozo de benefício",
                "Perda da qualidade de segurado",
                "Outro",
            ],
            "accept_new_options": True,
        },
        {
            "id": "carencia_cumprida",
            "label": "Carência cumprida?",
            "widget": "select",
            "options": [
                "Sim",
                "Não",
                "Em discussão",
            ],
        },
        {
            "id": "tempo_contribuicao_detalhe",
            "label": "Resumo detalhado dos períodos (opcional)",
            "widget": "textarea",
            "height": 90,
        },
        {
            "id": "incapacidade_limitacao",
            "label": "Incapacidade/limitação funcional (se houver)",
            "widget": "textarea",
            "height": 90,
            "placeholder": "Descreva limitações e impacto no trabalho/vida diária.",
        },
    ],
    "Direito da Saude": [
        {
            "id": "natureza_relacao_juridica",
            "label": "Natureza da relação",
            "widget": "select",
            "options": NATUREZA_RELACAO_OPCOES["Direito da Saude"],
            "accept_new_options": True,
        },
        {
            "id": "reu_tipo_saude",
            "label": "Quem é o réu?",
            "widget": "select",
            "options": [
                "Plano de saúde",
                "Município",
                "Estado",
                "União",
                "Hospital/Clínica",
                "Outro",
            ],
            "accept_new_options": True,
        },
        {
            "id": "tratamento_medicamento",
            "label": "Tratamento/medicamento/procedimento",
            "widget": "textarea",
            "height": 90,
            "placeholder": "Nome, dose, periodicidade, duração (se souber).",
        },
        {
            "id": "urgencia_laudo",
            "label": "Urgência e indicação médica (conforme laudo/relatório)",
            "widget": "textarea",
            "height": 90,
            "placeholder": "Resuma o que consta no laudo/relatório, sem inventar.",
        },
        {
            "id": "negativa_motivo",
            "label": "Negativa e motivo alegado (se houver)",
            "widget": "textarea",
            "height": 90,
            "placeholder": "Ex.: carência, 'fora do rol', 'experimental', falta de estoque, etc.",
        },
        {
            "id": "prazo_cumprimento",
            "label": "Prazo desejado para cumprimento (opcional)",
            "widget": "text",
            "placeholder": "Ex.: 24h, 48h, 5 dias",
        },
        {
            "id": "astreintes",
            "label": "Multa diária (astreintes) sugerida (opcional)",
            "widget": "text",
            "placeholder": "Ex.: R$ 1.000,00/dia",
        },
    ],
    "Outro": [
        {
            "id": "contexto_setorial",
            "label": "Contexto técnico/setorial da causa",
            "widget": "textarea",
            "height": 90,
            "placeholder": "Explique o contexto especializado do caso.",
        },
        {
            "id": "objeto_principal",
            "label": "Objeto principal da pretensão",
            "widget": "text",
            "placeholder": "Ex.: declaração de nulidade de cláusula X",
        },
        {
            "id": "riscos_sensiveis",
            "label": "Riscos/pontos sensíveis",
            "widget": "textarea",
            "height": 90,
            "placeholder": "Aspectos que exigem cuidado na redação.",
        },
    ],
}


 # Retorna apenas os caracteres numéricos de um texto.
def _somente_digitos(valor: str) -> str:
    return re.sub(r"\D", "", valor or "")


 # Formata dígitos em padrão de CPF, inclusive durante digitação parcial.
def _formatar_cpf(digitos: str) -> str:
    if len(digitos) <= 3:
        return digitos
    if len(digitos) <= 6:
        return f"{digitos[:3]}.{digitos[3:]}"
    if len(digitos) <= 9:
        return f"{digitos[:3]}.{digitos[3:6]}.{digitos[6:]}"
    return f"{digitos[:3]}.{digitos[3:6]}.{digitos[6:9]}-{digitos[9:11]}"


 # Formata dígitos em padrão de CNPJ, inclusive durante digitação parcial.
def _formatar_cnpj(digitos: str) -> str:
    if len(digitos) <= 2:
        return digitos
    if len(digitos) <= 5:
        return f"{digitos[:2]}.{digitos[2:]}"
    if len(digitos) <= 8:
        return f"{digitos[:2]}.{digitos[2:5]}.{digitos[5:]}"
    if len(digitos) <= 12:
        return f"{digitos[:2]}.{digitos[2:5]}.{digitos[5:8]}/{digitos[8:]}"
    return f"{digitos[:2]}.{digitos[2:5]}.{digitos[5:8]}/{digitos[8:12]}-{digitos[12:14]}"


 # Decide entre máscara de CPF ou CNPJ conforme a quantidade de dígitos.
def _formatar_cpf_cnpj(valor: str) -> str:
    digitos = _somente_digitos(valor)
    if len(digitos) <= 11:
        return _formatar_cpf(digitos[:11])
    return _formatar_cnpj(digitos[:14])


 # Converte uma sequência numérica para formato monetário brasileiro.
def _formatar_moeda_br(valor: str) -> str:
    digitos = _somente_digitos(valor)
    if not digitos:
        return ""
    centavos = int(digitos)
    inteiro = centavos // 100
    resto = centavos % 100
    inteiro_formatado = f"{inteiro:,}".replace(",", ".")
    return f"R$ {inteiro_formatado},{resto:02d}"


 # Converte texto multilinha em lista, removendo marcadores e linhas vazias.
def _linhas_para_lista(texto: str) -> list[str]:
    itens: list[str] = []
    for linha in (texto or "").splitlines():
        item = linha.strip()
        if item.startswith("-"):
            item = item[1:].strip()
        if item:
            itens.append(item)
    return itens


 # Mescla listas de textos sem duplicar itens (comparação case-insensitive).
def _mesclar_itens(*colecoes: list[str]) -> list[str]:
    resultado: list[str] = []
    vistos: set[str] = set()

    for colecao in colecoes:
        for item in colecao:
            texto = (item or "").strip()
            if not texto:
                continue
            chave = texto.casefold()
            if chave in vistos:
                continue
            vistos.add(chave)
            resultado.append(texto)

    return resultado


 # Retorna o modo de preenchimento atual do formulario.
def _modo_preenchimento() -> str:
    valor = str(st.session_state.get("modo_preenchimento", "Essencial")).strip()
    return valor if valor in MODO_PREENCHIMENTO_OPCOES else "Essencial"


 # Indica se o modo essencial (enxuto) esta ativo.
def _modo_essencial_ativo() -> bool:
    return _modo_preenchimento() == "Essencial"


 # Lista de pedidos base exibidos no formulario, sem os que ja possuem checkbox dedicado.
def _pedidos_base_exibicao() -> list[str]:
    return [pedido for pedido in PEDIDOS_BASE if pedido not in PEDIDOS_PARAMETROS_FINAIS]


 # Adiciona automaticamente pedidos derivados dos parametros finais.
def _incluir_pedidos_dos_parametros_finais(pedidos: list[str]) -> list[str]:
    extras: list[str] = []
    if st.session_state.get("tem_tutela_urgencia", False):
        extras.append("Tutela de urgência")
    if st.session_state.get("tem_gratuidade", False):
        extras.append("Justiça gratuita")
    return _mesclar_itens(pedidos, extras)


 # Gera um identificador simples e estável para uso em chaves de estado.
def _slug(valor: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", (valor or "").lower()).strip("_")


 # Monta a chave de session_state para um campo dinâmico de área.
def _chave_campo_area(area: str, campo_id: str) -> str:
    return f"area_{_slug(area)}_{_slug(campo_id)}"


 # Resolve o nome da área exibida para a chave interna usada no dicionário de campos.
def _resolver_area_campos(area: str) -> str:
    area_normalizada = (area or "").strip()
    area_mapeada = ALIAS_AREA_CAMPOS.get(area_normalizada, area_normalizada)
    return area_mapeada if area_mapeada in CAMPOS_POR_AREA else "Outro"


 # Clona estruturas mutáveis para evitar referência compartilhada no snapshot.
def _clonar_valor_snapshot(valor: Any) -> Any:
    if isinstance(valor, list):
        return valor.copy()
    if isinstance(valor, dict):
        return valor.copy()
    if isinstance(valor, set):
        return set(valor)
    if isinstance(valor, tuple):
        return tuple(valor)
    return valor


 # Lista todas as chaves do formulário que devem participar da persistência.
def _listar_todas_chaves_formulario() -> list[str]:
    chaves = list(CHAVES_FORMULARIO_BASE)
    for area, campos in CAMPOS_POR_AREA.items():
        for campo in campos:
            campo_id = str(campo.get("id", "")).strip()
            if campo_id:
                chaves.append(_chave_campo_area(area, campo_id))

    dedup: list[str] = []
    vistos: set[str] = set()
    for chave in chaves:
        if chave in vistos:
            continue
        vistos.add(chave)
        dedup.append(chave)
    return dedup


 # Salva um snapshot dos campos do formulário no session_state.
def _salvar_snapshot_formulario() -> None:
    snapshot = st.session_state.get("_form_snapshot", {})
    if not isinstance(snapshot, dict):
        snapshot = {}

    for chave in _listar_todas_chaves_formulario():
        if chave in st.session_state:
            snapshot[chave] = _clonar_valor_snapshot(st.session_state[chave])

    st.session_state["_form_snapshot"] = snapshot


 # Restaura chaves do snapshot quando elas não existem no ciclo atual de renderização.
def _restaurar_snapshot_formulario() -> None:
    snapshot = st.session_state.get("_form_snapshot", {})
    if not isinstance(snapshot, dict):
        return

    for chave, valor in snapshot.items():
        if chave not in st.session_state:
            st.session_state[chave] = _clonar_valor_snapshot(valor)


 # Renderiza um campo dinâmico com base na configuração do tipo de widget.
def _renderizar_campo_area(area: str, campo: dict[str, Any]) -> None:
    campo_id = str(campo.get("id", "")).strip()
    if not campo_id:
        return

    chave = _chave_campo_area(area, campo_id)
    rotulo = str(campo.get("label", campo_id))
    tipo_widget = str(campo.get("widget", "text")).strip().lower()
    placeholder = str(campo.get("placeholder", ""))
    ajuda = campo.get("help")
    altura = int(campo.get("height", 90))

    if tipo_widget == "textarea":
        st.text_area(rotulo, key=chave, height=altura, placeholder=placeholder, help=ajuda)
        return

    if tipo_widget == "select":
        opcoes = [str(item) for item in campo.get("options", [])]
        if not opcoes:
            st.selectbox(rotulo, ["[PREENCHER]"], key=chave, help=ajuda)
            return

        aceitar_novas_opcoes = bool(campo.get("accept_new_options", campo_id == "natureza_relacao_juridica"))
        if aceitar_novas_opcoes:
            kwargs_select: dict[str, Any] = {
                "label": rotulo,
                "options": opcoes,
                "key": chave,
                "help": ajuda,
                "placeholder": placeholder or "Selecione ou digite para buscar/criar...",
                "accept_new_options": True,
            }
            if chave not in st.session_state:
                kwargs_select["index"] = None
            st.selectbox(**kwargs_select)
            return

        st.selectbox(rotulo, opcoes, key=chave, help=ajuda)
        return

    if tipo_widget == "multiselect":
        opcoes = [str(item) for item in campo.get("options", [])]
        st.multiselect(rotulo, opcoes, key=chave, help=ajuda)
        return

    if tipo_widget == "checkbox":
        st.checkbox(rotulo, key=chave, help=ajuda)
        return

    st.text_input(rotulo, key=chave, placeholder=placeholder, help=ajuda)


 # Indica se o campo deve ocupar a linha inteira no bloco da area.
def _campo_area_linha_inteira(campo: dict[str, Any]) -> bool:
    widget = str(campo.get("widget", "text")).strip().lower()
    if widget in {"textarea", "multiselect"}:
        return True
    return bool(campo.get("full_width", False))


 # Identifica campos dinamicos opcionais com base no rotulo.
def _campo_area_eh_opcional(campo: dict[str, Any]) -> bool:
    rotulo = str(campo.get("label", "")).casefold()
    return "(opcional" in rotulo


 # Renderiza uma lista de campos em linhas de duas colunas.
def _renderizar_campos_area_em_duas_colunas(area_campos: str, campos: list[dict[str, Any]]) -> None:
    if not campos:
        return

    for idx in range(0, len(campos), 2):
        col1, col2 = st.columns(2)
        with col1:
            _renderizar_campo_area(area_campos, campos[idx])
        if idx + 1 < len(campos):
            with col2:
                _renderizar_campo_area(area_campos, campos[idx + 1])


 # Renderiza todos os campos específicos da área jurídica selecionada.
def _renderizar_bloco_area(area: str) -> None:
    area_campos = _resolver_area_campos(area)
    campos = CAMPOS_POR_AREA.get(area_campos, CAMPOS_POR_AREA["Outro"])
    st.caption(f"Campos especificos para a area selecionada: {area}")

    if _modo_essencial_ativo():
        campos = [campo for campo in campos if not _campo_area_eh_opcional(campo)]
        st.caption("Modo Essencial ativo: campos opcionais da area foram ocultados.")

    buffer_duas_colunas: list[dict[str, Any]] = []
    for campo in campos:
        if _campo_area_linha_inteira(campo):
            _renderizar_campos_area_em_duas_colunas(area_campos, buffer_duas_colunas)
            buffer_duas_colunas = []
            _renderizar_campo_area(area_campos, campo)
            continue

        buffer_duas_colunas.append(campo)

    _renderizar_campos_area_em_duas_colunas(area_campos, buffer_duas_colunas)


 # Coleta apenas os campos específicos da área que foram efetivamente preenchidos.
def _coletar_campos_area_especificos(area: str) -> dict[str, Any]:
    area_campos = _resolver_area_campos(area)
    campos = CAMPOS_POR_AREA.get(area_campos, CAMPOS_POR_AREA["Outro"])
    rotulos: dict[str, str] = {}
    valores: dict[str, Any] = {}

    for campo in campos:
        campo_id = str(campo.get("id", "")).strip()
        if not campo_id:
            continue

        rotulos[campo_id] = str(campo.get("label", campo_id))
        chave = _chave_campo_area(area_campos, campo_id)
        valor = st.session_state.get(chave)

        if isinstance(valor, str):
            valor = valor.strip()
            if not valor:
                continue
        elif isinstance(valor, list):
            valor = [str(item).strip() for item in valor if str(item).strip()]
            if not valor:
                continue
        elif isinstance(valor, bool):
            if not valor:
                continue
        elif valor is None:
            continue

        valores[campo_id] = valor

    return {
        "area": area,
        "rotulos": rotulos,
        "valores": valores,
    }


 # Remove caracteres inválidos para nome de arquivo no Windows.
def _sanitizar_nome_arquivo(texto: str) -> str:
    nome = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "", texto or "")
    nome = re.sub(r"\s+", " ", nome).strip(" .")
    return nome


 # Gera o nome padrão do DOCX com base no nome do autor.
def _nome_arquivo_docx(autor_nome: str) -> str:
    nome_autor = _sanitizar_nome_arquivo(autor_nome)
    if not nome_autor:
        nome_autor = "[NOME DO AUTOR]"
    return f"Petição Inicial - {nome_autor}.docx"


 # Gera o nome padrão do PDF com base no nome do autor.
def _nome_arquivo_pdf(autor_nome: str) -> str:
    nome_autor = _sanitizar_nome_arquivo(autor_nome)
    if not nome_autor:
        nome_autor = "[NOME DO AUTOR]"
    return f"Petição Inicial - {nome_autor}.pdf"


 # Aplica máscara de CPF/CNPJ em um campo de documento.
def _aplicar_mascara_documento(campo: str) -> None:
    st.session_state[campo] = _formatar_cpf_cnpj(st.session_state.get(campo, ""))


 # Aplica máscara de moeda brasileira em um campo monetário.
def _aplicar_mascara_moeda(campo: str) -> None:
    st.session_state[campo] = _formatar_moeda_br(st.session_state.get(campo, ""))


 # Aplica máscara de CEP em um campo.
def _aplicar_mascara_cep(campo: str) -> None:
    st.session_state[campo] = _formatar_cep_br(st.session_state.get(campo, ""))


 # Aplica todas as máscaras necessárias antes da geração da peça.
def _aplicar_mascaras_formulario() -> None:
    _aplicar_mascara_documento("autor_doc")
    _aplicar_mascara_documento("reu_doc")
    _aplicar_mascara_cep("autor_cep")
    _aplicar_mascara_cep("reu_cep")
    _aplicar_mascara_moeda("valor_causa")


 # Sugere o tipo de pessoa mais comum por área para facilitar o preenchimento.
def _sugerir_tipo_pessoa(area_direito: str, papel: str) -> str:
    area = (area_direito or "").strip()
    papel_norm = (papel or "").strip().lower()
    if area == "Previdenciário":
        return "Pessoa Física" if papel_norm == "autor" else "Pessoa Jurídica"
    if area == "Direito da Saúde":
        return "Pessoa Física" if papel_norm == "autor" else "Pessoa Jurídica"
    return "Pessoa Física"


 # Renderiza o seletor PF/PJ com sugestão inicial por área e parte processual.
def _renderizar_tipo_pessoa_parte(papel: str, area_direito: str) -> None:
    key = f"{papel}_tipo_pessoa"
    kwargs: dict[str, Any] = {
        "label": "Selecioene o tipo de pessoa",
        "options": TIPOS_PESSOA_OPCOES,
        "key": key,
        "horizontal": True,
    }
    if key not in st.session_state:
        sugestao = _sugerir_tipo_pessoa(area_direito, papel)
        kwargs["index"] = TIPOS_PESSOA_OPCOES.index(sugestao) if sugestao in TIPOS_PESSOA_OPCOES else 0
    st.radio(**kwargs)


 # Retorna texto limpo de um campo no session_state.
def _texto_campo(chave: str) -> str:
    return str(st.session_state.get(chave, "")).strip()


 # Monta a qualificação textual conforme PF/PJ para cada parte.
def _montar_qualificacao_parte(papel: str, tipo_pessoa: str) -> str:
    prefixo = (papel or "").strip().lower()
    base_extra = _texto_campo(f"{prefixo}_qualificacao")
    partes: list[str] = []

    if tipo_pessoa == "Pessoa Jurídica":
        natureza_juridica = _texto_campo(f"{prefixo}_natureza_juridica")
        representante_legal = _texto_campo(f"{prefixo}_representante_legal")
        if natureza_juridica:
            partes.append(f"Natureza jurídica: {natureza_juridica}")
        if representante_legal:
            partes.append(f"Representante legal: {representante_legal}")
    else:
        nacionalidade = _texto_campo(f"{prefixo}_nacionalidade")
        estado_civil = _texto_campo(f"{prefixo}_estado_civil")
        profissao = _texto_campo(f"{prefixo}_profissao")
        if nacionalidade:
            partes.append(f"Nacionalidade: {nacionalidade}")
        if estado_civil:
            partes.append(f"Estado civil: {estado_civil}")
        if profissao:
            partes.append(f"Profissão: {profissao}")

    if base_extra:
        partes.append(base_extra)

    return "; ".join(partes)


 # Consolida os dados de uma parte (autor/réu) com estrutura PF/PJ.
def _coletar_dados_parte(papel: str) -> dict[str, str]:
    prefixo = (papel or "").strip().lower()
    tipo_pessoa = _texto_campo(f"{prefixo}_tipo_pessoa") or "Pessoa Física"
    qualificacao = _montar_qualificacao_parte(prefixo, tipo_pessoa)

    return {
        "tipo_pessoa": tipo_pessoa,
        "nome": _texto_campo(f"{prefixo}_nome"),
        "documento": _formatar_cpf_cnpj(_texto_campo(f"{prefixo}_doc")),
        "cep": _formatar_cep_br(_texto_campo(f"{prefixo}_cep")),
        "endereco": _texto_campo(f"{prefixo}_end"),
        "qualificacao": qualificacao,
        "nacionalidade": _texto_campo(f"{prefixo}_nacionalidade"),
        "estado_civil": _texto_campo(f"{prefixo}_estado_civil"),
        "profissao": _texto_campo(f"{prefixo}_profissao"),
        "natureza_juridica": _texto_campo(f"{prefixo}_natureza_juridica"),
        "representante_legal": _texto_campo(f"{prefixo}_representante_legal"),
    }


 # Formata CEP brasileiro no padrão 00000-000.
def _formatar_cep_br(valor: str) -> str:
    digitos = _somente_digitos(valor)
    if len(digitos) <= 5:
        return digitos
    return f"{digitos[:5]}-{digitos[5:8]}"


 # Limita texto de referencia para manter o prompt em tamanho controlado.
def _limitar_texto_modelo_referencia(texto: str, limite: int = LIMITE_CARACTERES_MODELO_REFERENCIA) -> tuple[str, bool]:
    conteudo = str(texto or "").replace("\r\n", "\n").replace("\r", "\n")
    conteudo = re.sub(r"\n{3,}", "\n\n", conteudo).strip()
    if len(conteudo) <= limite:
        return conteudo, False
    trecho = conteudo[:limite].rstrip()
    aviso = "\n...[MODELO TRUNCADO PARA CABER NO PROMPT]..."
    return f"{trecho}{aviso}", True


 # Decodifica arquivo textual usando codificacoes comuns.
def _decodificar_texto_arquivo(conteudo: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return conteudo.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Nao foi possivel ler o arquivo textual. Use .txt/.md em UTF-8.")


 # Extrai texto de arquivo DOCX para uso como referencia.
def _extrair_texto_docx_referencia(conteudo: bytes) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise ValueError("Leitura de .docx indisponivel no ambiente.") from exc

    try:
        documento = Document(io.BytesIO(conteudo))
    except Exception as exc:
        raise ValueError("Nao foi possivel ler o arquivo .docx enviado.") from exc

    blocos: list[str] = []
    for paragrafo in documento.paragraphs:
        texto = str(paragrafo.text or "").strip()
        if texto:
            blocos.append(texto)

    for tabela in documento.tables:
        for linha in tabela.rows:
            celulas = [str(celula.text or "").strip() for celula in linha.cells]
            celulas = [celula for celula in celulas if celula]
            if celulas:
                blocos.append(" | ".join(celulas))

    return "\n".join(blocos).strip()


 # Extrai texto do arquivo de modelo conforme extensao permitida.
def _extrair_texto_arquivo_modelo(nome_arquivo: str, conteudo: bytes) -> str:
    extensao = os.path.splitext(str(nome_arquivo or "").lower())[1]
    if extensao in {".txt", ".md", ".markdown"}:
        return _decodificar_texto_arquivo(conteudo)
    if extensao == ".docx":
        return _extrair_texto_docx_referencia(conteudo)
    raise ValueError("Formato nao suportado. Envie .txt, .md ou .docx.")


 # Processa o upload do modelo de referencia e persiste texto extraido no estado.
def _processar_modelo_referencia(uploaded_file: Any) -> None:
    nome_key = "modelo_referencia_nome"
    texto_key = "modelo_referencia_texto"
    trunc_key = "modelo_referencia_truncado"
    assinatura_key = "_modelo_referencia_assinatura"
    erro_key = "_modelo_referencia_erro"

    if uploaded_file is None:
        st.session_state.pop(assinatura_key, None)
        st.session_state.pop(erro_key, None)
        st.session_state.pop(nome_key, None)
        st.session_state.pop(texto_key, None)
        st.session_state.pop(trunc_key, None)
        return

    nome_arquivo = str(getattr(uploaded_file, "name", "")).strip()
    conteudo = bytes(uploaded_file.getvalue() or b"")

    assinatura_raw = f"{nome_arquivo}|{len(conteudo)}".encode("utf-8") + conteudo
    assinatura = hashlib.sha1(assinatura_raw).hexdigest()
    if assinatura == st.session_state.get(assinatura_key):
        return

    try:
        texto_extraido = _extrair_texto_arquivo_modelo(nome_arquivo, conteudo)
    except ValueError as exc:
        st.session_state[erro_key] = str(exc)
        st.session_state.pop(nome_key, None)
        st.session_state.pop(texto_key, None)
        st.session_state.pop(trunc_key, None)
        st.session_state[assinatura_key] = assinatura
        return
    except Exception:
        st.session_state[erro_key] = "Erro inesperado ao processar o modelo de referencia."
        st.session_state.pop(nome_key, None)
        st.session_state.pop(texto_key, None)
        st.session_state.pop(trunc_key, None)
        st.session_state[assinatura_key] = assinatura
        return

    texto_limitado, truncado = _limitar_texto_modelo_referencia(texto_extraido)
    if not texto_limitado:
        st.session_state[erro_key] = "Nao foi possivel extrair texto util do arquivo enviado."
        st.session_state.pop(nome_key, None)
        st.session_state.pop(texto_key, None)
        st.session_state.pop(trunc_key, None)
        st.session_state[assinatura_key] = assinatura
        return

    st.session_state[nome_key] = nome_arquivo or "[SEM NOME]"
    st.session_state[texto_key] = texto_limitado
    st.session_state[trunc_key] = truncado
    st.session_state.pop(erro_key, None)
    st.session_state[assinatura_key] = assinatura


 # Extrai nome de representante legal/sócio principal do retorno da BrasilAPI.
def _extrair_representante_brasilapi(dados_api: dict[str, Any]) -> str:
    qsa = dados_api.get("qsa", [])
    if not isinstance(qsa, list):
        return ""

    for socio in qsa:
        if not isinstance(socio, dict):
            continue
        nome_rep = str(socio.get("nome_representante_legal", "")).strip()
        if nome_rep:
            return nome_rep

    for socio in qsa:
        if not isinstance(socio, dict):
            continue
        nome_socio = str(socio.get("nome_socio", "")).strip()
        if nome_socio:
            return nome_socio

    return ""


 # Monta endereço textual a partir da resposta da BrasilAPI.
def _montar_endereco_pj_brasilapi(dados_api: dict[str, Any]) -> str:
    tipo_logradouro = str(dados_api.get("descricao_tipo_de_logradouro", "")).strip()
    logradouro = str(dados_api.get("logradouro", "")).strip()
    numero = str(dados_api.get("numero", "")).strip() or "S/N"
    complemento = str(dados_api.get("complemento", "")).strip()
    bairro = str(dados_api.get("bairro", "")).strip()
    municipio = str(dados_api.get("municipio", "")).strip()
    uf = str(dados_api.get("uf", "")).strip()
    cep = _formatar_cep_br(str(dados_api.get("cep", "")))

    linha_logradouro = " ".join(item for item in [tipo_logradouro, logradouro] if item).strip()
    if linha_logradouro:
        linha_logradouro = f"{linha_logradouro}, {numero}"
    else:
        linha_logradouro = ""

    partes = [linha_logradouro, complemento, bairro]

    cidade_uf = " / ".join(item for item in [municipio, uf] if item).strip(" /")
    if cidade_uf:
        partes.append(cidade_uf)
    if cep:
        partes.append(f"CEP {cep}")

    partes_validas = [parte for parte in partes if parte]
    return " - ".join(partes_validas)


 # Consulta dados públicos de CNPJ na BrasilAPI.
@st.cache_data(ttl=3600, show_spinner=False)
def _consultar_cnpj_brasilapi(cnpj_digitos: str) -> dict[str, Any]:
    digitos = _somente_digitos(cnpj_digitos)
    if len(digitos) != 14:
        raise ValueError("Informe um CNPJ com 14 dígitos para consulta.")

    url = f"https://brasilapi.com.br/api/cnpj/v1/{digitos}"
    req = urlrequest.Request(url, headers={"User-Agent": "streamlit-app/1.0"})

    try:
        with urlrequest.urlopen(req, timeout=12) as resp:
            conteudo = resp.read().decode("utf-8")
    except urlerror.HTTPError as exc:
        if exc.code == 404:
            raise ValueError("CNPJ não encontrado na BrasilAPI.") from exc
        raise ValueError(f"Falha ao consultar CNPJ na BrasilAPI (HTTP {exc.code}).") from exc
    except urlerror.URLError as exc:
        raise ValueError("Não foi possível conectar à BrasilAPI.") from exc
    except TimeoutError as exc:
        raise ValueError("Tempo esgotado na consulta do CNPJ.") from exc

    try:
        dados_api = json.loads(conteudo)
    except json.JSONDecodeError as exc:
        raise ValueError("Resposta inválida da BrasilAPI.") from exc

    if not isinstance(dados_api, dict):
        raise ValueError("Resposta inválida da BrasilAPI.")
    return dados_api


 # Consulta dados públicos de CEP na BrasilAPI.
@st.cache_data(ttl=3600, show_spinner=False)
def _consultar_cep_brasilapi(cep_digitos: str) -> dict[str, Any]:
    digitos = _somente_digitos(cep_digitos)
    if len(digitos) != 8:
        raise ValueError("Informe um CEP com 8 dígitos para consulta.")

    url = f"https://brasilapi.com.br/api/cep/v1/{digitos}"
    req = urlrequest.Request(url, headers={"User-Agent": "streamlit-app/1.0"})

    try:
        with urlrequest.urlopen(req, timeout=12) as resp:
            conteudo = resp.read().decode("utf-8")
    except urlerror.HTTPError as exc:
        if exc.code == 404:
            raise ValueError("CEP não encontrado na BrasilAPI.") from exc
        raise ValueError(f"Falha ao consultar CEP na BrasilAPI (HTTP {exc.code}).") from exc
    except urlerror.URLError as exc:
        raise ValueError("Não foi possível conectar à BrasilAPI para consulta de CEP.") from exc
    except TimeoutError as exc:
        raise ValueError("Tempo esgotado na consulta do CEP.") from exc

    try:
        dados_api = json.loads(conteudo)
    except json.JSONDecodeError as exc:
        raise ValueError("Resposta inválida da BrasilAPI para CEP.") from exc

    if not isinstance(dados_api, dict):
        raise ValueError("Resposta inválida da BrasilAPI para CEP.")
    return dados_api


 # Monta endereço a partir do retorno da BrasilAPI de CEP.
def _montar_endereco_cep_brasilapi(dados_api: dict[str, Any]) -> str:
    rua = str(dados_api.get("street", "")).strip()
    bairro = str(dados_api.get("neighborhood", "")).strip()
    cidade = str(dados_api.get("city", "")).strip()
    uf = str(dados_api.get("state", "")).strip()
    cep = _formatar_cep_br(str(dados_api.get("cep", "")))

    partes = [rua, bairro]
    cidade_uf = " / ".join(item for item in [cidade, uf] if item).strip(" /")
    if cidade_uf:
        partes.append(cidade_uf)
    if cep:
        partes.append(f"CEP {cep}")

    partes_validas = [item for item in partes if item]
    return " - ".join(partes_validas)


 # Preenche campos de parte PJ com base no CNPJ consultado na BrasilAPI.
def _preencher_parte_com_cnpj(papel: str) -> None:
    prefixo = (papel or "").strip().lower()
    feedback_key = f"_{prefixo}_cnpj_feedback"

    cnpj_digitos = _somente_digitos(_texto_campo(f"{prefixo}_doc"))
    if len(cnpj_digitos) != 14:
        st.session_state[feedback_key] = ("error", "Informe um CNPJ válido (14 dígitos) antes de buscar.")
        return

    try:
        dados_api = _consultar_cnpj_brasilapi(cnpj_digitos)
    except ValueError as exc:
        st.session_state[feedback_key] = ("error", str(exc))
        return
    except Exception:
        st.session_state[feedback_key] = ("error", "Erro inesperado ao consultar CNPJ.")
        return

    razao_social = str(dados_api.get("razao_social", "")).strip()
    nome_fantasia = str(dados_api.get("nome_fantasia", "")).strip()
    natureza_juridica = str(dados_api.get("natureza_juridica", "")).strip()
    endereco = _montar_endereco_pj_brasilapi(dados_api)
    cep_cnpj = _formatar_cep_br(str(dados_api.get("cep", "")))
    representante_legal = _extrair_representante_brasilapi(dados_api)
    situacao = str(dados_api.get("descricao_situacao_cadastral", "")).strip()
    email = str(dados_api.get("email", "")).strip()
    telefone = _somente_digitos(str(dados_api.get("ddd_telefone_1", "")))

    st.session_state[f"{prefixo}_tipo_pessoa"] = "Pessoa Jurídica"
    if razao_social:
        st.session_state[f"{prefixo}_nome"] = razao_social
    elif nome_fantasia:
        st.session_state[f"{prefixo}_nome"] = nome_fantasia

    if natureza_juridica:
        st.session_state[f"{prefixo}_natureza_juridica"] = natureza_juridica
    if representante_legal:
        st.session_state[f"{prefixo}_representante_legal"] = representante_legal
    if cep_cnpj:
        st.session_state[f"{prefixo}_cep"] = cep_cnpj
    if endereco:
        st.session_state[f"{prefixo}_end"] = endereco

    extras: list[str] = []
    if situacao:
        extras.append(f"Situação cadastral: {situacao}")
    if telefone:
        extras.append(f"Telefone: {telefone}")
    if email:
        extras.append(f"E-mail: {email}")
    if extras and not _texto_campo(f"{prefixo}_qualificacao"):
        st.session_state[f"{prefixo}_qualificacao"] = "; ".join(extras)

    st.session_state[feedback_key] = ("success", "Dados da pessoa jurídica preenchidos.")


 # Preenche endereço da parte com base no CEP consultado.
def _preencher_endereco_por_cep(papel: str) -> None:
    prefixo = (papel or "").strip().lower()
    feedback_key = f"_{prefixo}_cep_feedback"

    cep_digitos = _somente_digitos(_texto_campo(f"{prefixo}_cep"))
    if len(cep_digitos) != 8:
        st.session_state[feedback_key] = ("error", "Informe um CEP válido (8 dígitos) antes de buscar.")
        return

    try:
        dados_api = _consultar_cep_brasilapi(cep_digitos)
    except ValueError as exc:
        st.session_state[feedback_key] = ("error", str(exc))
        return
    except Exception:
        st.session_state[feedback_key] = ("error", "Erro inesperado ao consultar CEP.")
        return

    endereco = _montar_endereco_cep_brasilapi(dados_api)
    cep_fmt = _formatar_cep_br(str(dados_api.get("cep", cep_digitos)))
    if cep_fmt:
        st.session_state[f"{prefixo}_cep"] = cep_fmt

    if endereco:
        st.session_state[f"{prefixo}_end"] = endereco
        st.session_state[feedback_key] = ("success", "Endereço preenchido via CEP. Revise número/complemento se necessário.")
    else:
        st.session_state[feedback_key] = ("error", "CEP encontrado, mas sem dados suficientes para montar endereço.")


 # Mostra feedback de consulta CNPJ e remove a mensagem após exibir.
def _exibir_feedback_cnpj(papel: str) -> None:
    prefixo = (papel or "").strip().lower()
    feedback_key = f"_{prefixo}_cnpj_feedback"
    feedback = st.session_state.pop(feedback_key, None)
    if not isinstance(feedback, tuple) or len(feedback) != 2:
        return

    tipo, mensagem = feedback
    texto = str(mensagem or "").strip()
    if not texto:
        return

    if str(tipo) == "success":
        st.success(texto)
    else:
        st.error(texto)


 # Mostra feedback de consulta de CEP e remove a mensagem após exibir.
def _exibir_feedback_cep(papel: str) -> None:
    prefixo = (papel or "").strip().lower()
    feedback_key = f"_{prefixo}_cep_feedback"
    feedback = st.session_state.pop(feedback_key, None)
    if not isinstance(feedback, tuple) or len(feedback) != 2:
        return

    tipo, mensagem = feedback
    texto = str(mensagem or "").strip()
    if not texto:
        return

    if str(tipo) == "success":
        st.success(texto)
    else:
        st.error(texto)


 # Consolida os dados do formulário no payload usado pelo prompt.
def _coletar_payload() -> dict[str, Any]:
    area_direito = st.session_state.get("area_direito", "Outro")
    campos_area_especificos = _coletar_campos_area_especificos(area_direito)
    valor_causa_fmt = _formatar_moeda_br(str(st.session_state.get("valor_causa", "")))

    pedidos_custom = _linhas_para_lista(st.session_state.get("pedidos_custom_raw", ""))
    pedidos_base_raw = st.session_state.get("pedidos_base", [])
    if not isinstance(pedidos_base_raw, list):
        pedidos_base_raw = []
    pedidos_base = [str(item).strip() for item in pedidos_base_raw if str(item).strip()]
    pedidos_base = _incluir_pedidos_dos_parametros_finais(pedidos_base)
    pedidos_lista_final = _mesclar_itens(pedidos_base, pedidos_custom)

    fundamentos_legais = _linhas_para_lista(st.session_state.get("fundamentos_legais_raw", ""))
    provas_documentos_raw = _linhas_para_lista(st.session_state.get("provas_raw", ""))
    provas_sugeridas_raw = st.session_state.get("provas_sugeridas", [])
    if not isinstance(provas_sugeridas_raw, list):
        provas_sugeridas_raw = []
    provas_sugeridas = [str(item).strip() for item in provas_sugeridas_raw if str(item).strip()]
    provas_documentos = _mesclar_itens(provas_sugeridas, provas_documentos_raw)
    cronologia = _linhas_para_lista(st.session_state.get("cronologia_raw", ""))
    partes_adicionais = _linhas_para_lista(st.session_state.get("partes_adicionais_raw", ""))
    secoes_extras = _linhas_para_lista(st.session_state.get("secoes_extras_raw", ""))

    temas_custom = _linhas_para_lista(st.session_state.get("temas_custom_raw", ""))
    temas_comuns = st.session_state.get("temas_comuns", [])
    temas_juridicos = _mesclar_itens(temas_comuns, temas_custom)

    autor = _coletar_dados_parte("autor")
    reu = _coletar_dados_parte("reu")

    advogado = {
        "nome": st.session_state.get("advogado_nome", ""),
        "oab_uf": str(st.session_state.get("advogado_oab_uf", "")).strip().upper(),
        "oab_num": st.session_state.get("advogado_oab_num", ""),
    }
    modelo_referencia = {
        "nome_arquivo": st.session_state.get("modelo_referencia_nome", ""),
        "texto": st.session_state.get("modelo_referencia_texto", ""),
        "conteudo_truncado": bool(st.session_state.get("modelo_referencia_truncado", False)),
    }

    dados = {
        "contexto_processual": {
            "area_direito": area_direito,
            "tipo_acao": st.session_state.get("tipo_acao", ""),
            "rito": st.session_state.get("rito", ""),
            "comarca_uf": st.session_state.get("comarca_uf", ""),
            "foro_vara": st.session_state.get("foro_vara", ""),
        },
        "campos_area_especificos": campos_area_especificos,
        "partes": {
            "autor": autor,
            "reu": reu,
            "partes_adicionais": partes_adicionais,
        },
        "narrativa": {
            "fatos": st.session_state.get("fatos", ""),
            "cronologia": cronologia,
            "provas_sugeridas": provas_sugeridas,
            "provas_documentos": provas_documentos,
        },
        "fundamentacao": {
            "teses_juridicas": st.session_state.get("teses_juridicas", ""),
            "temas_juridicos": temas_juridicos,
            "fundamentos_legais": fundamentos_legais,
        },
        "pedidos": pedidos_lista_final,
        "pedidos_detalhados": {
            "base_selecionados": pedidos_base,
            "personalizados": pedidos_custom,
            "lista_final": pedidos_lista_final,
        },
        "estrutura_peticao": {
            "secoes_sugeridas": st.session_state.get("secoes_sugeridas", []),
            "secoes_extras": secoes_extras,
            "nivel_detalhamento": st.session_state.get("nivel_detalhamento", "Padrao"),
        },
        "parametros_finais": {
            "valor_causa": valor_causa_fmt,
            "tutela_urgencia": st.session_state.get("tem_tutela_urgencia", False),
            "justica_gratuita": st.session_state.get("tem_gratuidade", False),
            "prioridade_tramitacao": st.session_state.get("tem_prioridade", False),
            "audiencia_conciliacao": st.session_state.get("quer_audiencia", True),
        },
        "observacoes_estrategicas": st.session_state.get("obs_estrategicas", ""),
        "advogado": advogado,
        "modelo_referencia": modelo_referencia,
        "autor": autor,
        "reu": reu,
        "tipo_acao": st.session_state.get("tipo_acao", ""),
        "fatos": st.session_state.get("fatos", ""),
        "valor_causa": valor_causa_fmt,
    }

    return dados


 # Verifica se um campo possui conteúdo válido, respeitando o tipo do valor.
def _campo_preenchido(chave: str) -> bool:
    valor = st.session_state.get(chave)
    if isinstance(valor, str):
        return bool(valor.strip())
    if isinstance(valor, list):
        return len(valor) > 0
    if isinstance(valor, bool):
        return valor
    return valor is not None


 # Retorna as linhas com conteúdo de um campo textual multilinha.
def _linhas_com_texto(chave: str) -> list[str]:
    return _linhas_para_lista(st.session_state.get(chave, ""))


 # Obtém o índice da etapa atual com proteção de limites.
def _obter_etapa_idx() -> int:
    if "etapa_idx" not in st.session_state:
        st.session_state.etapa_idx = 0
    idx = int(st.session_state.etapa_idx)
    return max(0, min(idx, len(ETAPAS_FLUXO) - 1))


 # Define o índice da etapa atual com proteção de limites.
def _definir_etapa_idx(idx: int) -> None:
    st.session_state.etapa_idx = max(0, min(int(idx), len(ETAPAS_FLUXO) - 1))


 # Mapeia os campos obrigatórios de cada etapa do fluxo.
def _campos_obrigatorios_da_etapa(etapa: str) -> list[tuple[str, str]]:
    obrigatorios: dict[str, list[tuple[str, str]]] = {
        "Contexto Processual": [
            ("tipo_acao", "Tipo da ação"),
            ("comarca_uf", "Comarca / UF"),
        ],
        "Partes": [
            ("autor_nome", "Nome do autor"),
            ("autor_doc", "CPF/CNPJ do autor"),
            ("autor_end", "Endereço do autor"),
            ("reu_nome", "Nome do réu"),
            ("reu_doc", "CPF/CNPJ do réu"),
            ("reu_end", "Endereço do réu"),
        ],
        "Fatos e Provas": [
            ("fatos", "Fatos principais"),
        ],
        "Fundamentação": [
            ("teses_juridicas", "Teses jurídicas"),
        ],
        "Finalização e Geração": [
            ("valor_causa", "Valor da causa"),
        ],
    }
    return obrigatorios.get(etapa, [])


 # Valida se os obrigatórios da etapa atual foram preenchidos.
def _validar_etapa(etapa: str) -> list[str]:
    faltantes: list[str] = []

    for chave, rotulo in _campos_obrigatorios_da_etapa(etapa):
        if not _campo_preenchido(chave):
            faltantes.append(rotulo)

    if etapa == "Pedidos":
        pedidos_base = st.session_state.get("pedidos_base", [])
        pedidos_custom = _linhas_com_texto("pedidos_custom_raw")
        if not pedidos_base and not pedidos_custom:
            faltantes.append("Selecione ao menos um pedido (base ou personalizado)")

    return faltantes


 # Valida os campos essenciais de todas as etapas antes de gerar a petição.
def _validar_essenciais_para_geracao() -> list[str]:
    etapas_relevantes = [
        "Contexto Processual",
        "Partes",
        "Fatos e Provas",
        "Fundamentação",
        "Pedidos",
        "Finalização e Geração",
    ]
    faltantes: list[str] = []
    for etapa in etapas_relevantes:
        faltantes.extend(_validar_etapa(etapa))

    dedup: list[str] = []
    vistos: set[str] = set()
    for item in faltantes:
        chave = item.casefold()
        if chave in vistos:
            continue
        vistos.add(chave)
        dedup.append(item)
    return dedup


 # Calcula o progresso considerando etapas concluídas e etapa atual válida.
def _calcular_progresso_preenchimento() -> float:
    total_etapas = len(ETAPAS_FLUXO)
    if total_etapas == 0:
        return 0.0

    etapa_idx_atual = _obter_etapa_idx()
    etapas_concluidas = etapa_idx_atual

    etapa_atual = ETAPAS_FLUXO[etapa_idx_atual]
    if not _validar_etapa(etapa_atual):
        etapas_concluidas += 1

    progresso = etapas_concluidas / total_etapas
    return max(0.0, min(progresso, 1.0))


 # Renderiza o painel lateral de acompanhamento visual do fluxo.
def _menu_fluxo_lateral() -> tuple[str, int]:
    with st.sidebar:
        st.markdown("### Painel do Caso")
        st.caption("Navegue por etapas para montar a petição.")

        st.selectbox(
            "Área do direito",
            AREAS_DIREITO,
            key="area_direito",
        )
        st.radio(
            "Modo de preenchimento",
            MODO_PREENCHIMENTO_OPCOES,
            key="modo_preenchimento",
            horizontal=True,
            help="Essencial mostra so os campos principais. Completo exibe todos os opcionais.",
        )
        if _modo_essencial_ativo():
            st.caption("Modo Essencial ativo.")

        etapa_idx_atual = _obter_etapa_idx()
        st.markdown("### Fluxo de Preenchimento")

        linhas_fluxo: list[str] = []
        for idx, etapa_nome in enumerate(ETAPAS_FLUXO):
            if idx < etapa_idx_atual:
                classe = "concluida"
                marcador = "✓"
            elif idx == etapa_idx_atual:
                classe = "atual"
                marcador = str(idx + 1)
            else:
                classe = "pendente"
                marcador = str(idx + 1)

            linhas_fluxo.append(
                (
                    f'<div class="fluxo-item {classe}">'
                    f'<span class="fluxo-badge">{html.escape(marcador)}</span>'
                    f'<span class="fluxo-text">{html.escape(etapa_nome)}</span>'
                    "</div>"
                )
            )

        st.markdown(f'<div class="fluxo-tracker">{"".join(linhas_fluxo)}</div>', unsafe_allow_html=True)
        st.caption("Apenas acompanhamento visual. Use os botões Voltar/Avançar para mudar de etapa.")

        progresso = _calcular_progresso_preenchimento()
        st.progress(progresso)
        st.caption(f"Progresso por etapas: {int(progresso * 100)}%")

    return ETAPAS_FLUXO[etapa_idx_atual], etapa_idx_atual


 # Injeta o tema visual escuro-dourado da aplicação.
def _aplicar_estilo_preto_dourado() -> None:
    st.markdown(
        """
        <style>
            .stApp {
                background:
                    radial-gradient(1200px 520px at -8% -8%, rgba(51, 96, 186, 0.34), rgba(9, 18, 39, 0) 58%),
                    radial-gradient(950px 420px at 108% 4%, rgba(214, 170, 71, 0.2), rgba(9, 18, 39, 0) 54%),
                    linear-gradient(180deg, #0c1730 0%, #08101f 100%);
                color: #fbfbfb;
                font-family: "Montserrat", "Trebuchet MS", sans-serif;
            }

            .main .block-container {
                max-width: 1220px;
                padding-top: 1rem;
                padding-bottom: 2.4rem;
            }

            [data-testid="stSidebar"] {
                background: linear-gradient(180deg, #0f1f3f 0%, #0a1428 100%);
                border-right: 1px solid rgba(216, 171, 73, 0.34);
            }

            [data-testid="stSidebar"] .block-container {
                padding-top: 1rem;
            }

            .fluxo-tracker {
                margin-top: 0.45rem;
                border: 1px solid rgba(96, 145, 228, 0.34);
                border-radius: 14px;
                padding: 0.45rem 0.4rem;
                background: #122245;
            }

            .fluxo-item {
                display: flex;
                align-items: center;
                gap: 0.5rem;
                margin: 3px 0;
                padding: 0.42rem 0.45rem;
                border-radius: 10px;
                border: 1px solid transparent;
                transition: none;
            }

            .fluxo-item.pendente {
                color: #d9e6ff;
                background: rgba(84, 130, 214, 0.14);
            }

            .fluxo-item.concluida {
                color: #dff0ff;
                border-color: rgba(116, 170, 244, 0.58);
                background: rgba(66, 120, 211, 0.24);
            }

            .fluxo-item.atual {
                color: #241a05;
                border-color: rgba(255, 230, 166, 0.45);
                background: linear-gradient(135deg, #d0a64b, #f6dea6);
                box-shadow: 0 4px 12px rgba(0, 0, 0, 0.24);
            }

            .fluxo-badge {
                width: 1.22rem;
                height: 1.22rem;
                border-radius: 999px;
                border: 1px solid rgba(108, 157, 236, 0.5);
                display: inline-flex;
                align-items: center;
                justify-content: center;
                font-size: 0.72rem;
                font-weight: 700;
                color: #ddeaff;
                background: rgba(70, 117, 205, 0.3);
                flex: 0 0 1.22rem;
            }

            .fluxo-item.concluida .fluxo-badge {
                border-color: rgba(132, 188, 255, 0.72);
                color: #eff7ff;
                background: rgba(80, 134, 225, 0.38);
            }

            .fluxo-item.atual .fluxo-badge {
                border-color: rgba(44, 32, 7, 0.35);
                color: #2a1f05;
                background: rgba(255, 244, 216, 0.5);
            }

            .fluxo-text {
                font-size: 0.9rem;
                font-weight: 600;
                line-height: 1.2;
            }

            h1, h2, h3, label, p, span {
                color: #fbfbfb !important;
            }

            div[data-testid="stCaptionContainer"] p {
                color: #b7c8e8 !important;
            }

            .main div[data-testid="stVerticalBlockBorderWrapper"],
            div[data-testid="stForm"] {
                background: linear-gradient(180deg, rgba(21, 36, 70, 0.96), rgba(10, 20, 40, 0.96));
                border: 1px solid rgba(216, 171, 73, 0.42);
                border-radius: 20px;
                padding: 1.3rem;
                backdrop-filter: blur(8px);
                box-shadow: 0 18px 34px rgba(0, 0, 0, 0.35);
            }

            .hero-shell {
                position: relative;
                overflow: hidden;
                border-radius: 20px;
                border: 1px solid rgba(216, 171, 73, 0.42);
                background: linear-gradient(135deg, rgba(24, 43, 82, 0.96), rgba(11, 23, 46, 0.97));
                padding: 1.15rem 1.25rem 1rem 1.25rem;
                margin-bottom: 0.85rem;
                box-shadow: 0 14px 28px rgba(0, 0, 0, 0.3);
            }

            .hero-shell::after {
                content: "";
                position: absolute;
                inset: -1px;
                background:
                    radial-gradient(circle at 88% 15%, rgba(246, 217, 139, 0.16), rgba(0, 0, 0, 0) 45%),
                    radial-gradient(circle at 14% 90%, rgba(96, 145, 228, 0.16), rgba(0, 0, 0, 0) 48%);
                pointer-events: none;
            }

            .hero-title {
                position: relative;
                z-index: 1;
                margin: 0;
                font-size: 2rem;
                font-weight: 750;
                letter-spacing: 0.2px;
                color: #f6d98b !important;
                line-height: 1.15;
            }

            .hero-subtitle {
                position: relative;
                z-index: 1;
                margin-top: 0.35rem;
                color: #d6e2fb;
                font-size: 0.95rem;
                max-width: 74ch;
            }

            .hero-chip-row {
                position: relative;
                z-index: 1;
                display: flex;
                flex-wrap: wrap;
                gap: 0.45rem;
                margin-top: 0.8rem;
            }

            .hero-chip {
                display: inline-flex;
                align-items: center;
                border-radius: 999px;
                border: 1px solid rgba(104, 153, 235, 0.56);
                background: rgba(63, 108, 194, 0.2);
                color: #deebff;
                font-size: 0.8rem;
                letter-spacing: 0.15px;
                padding: 0.24rem 0.62rem;
            }

            .hero-chip.status-ok {
                border-color: rgba(113, 168, 245, 0.78);
                background: rgba(67, 121, 218, 0.3);
                color: #e8f3ff;
            }

            .hero-chip.status-warn {
                border-color: rgba(232, 187, 92, 0.78);
                background: rgba(216, 171, 73, 0.22);
                color: #ffe5af;
            }

            .secao-titulo {
                display: flex;
                align-items: center;
                gap: 0.55rem;
                margin-top: 1.1rem;
                margin-bottom: 0.6rem;
                font-weight: 700;
                letter-spacing: 0.22px;
                color: #f6d98b;
                font-size: 1rem;
            }

            .secao-indice {
                width: 1.6rem;
                height: 1.6rem;
                border-radius: 999px;
                border: 1px solid rgba(216, 171, 73, 0.65);
                background: rgba(216, 171, 73, 0.16);
                display: inline-flex;
                align-items: center;
                justify-content: center;
                color: #ffe8b8;
                font-size: 0.77rem;
                font-weight: 700;
                flex: 0 0 1.6rem;
            }

            .secao-linha {
                flex: 1;
                height: 1px;
                background: linear-gradient(90deg, rgba(216, 171, 73, 0.55), rgba(216, 171, 73, 0.04));
            }

            div[data-baseweb="input"] > div,
            div[data-baseweb="textarea"] > div,
            div[data-baseweb="select"] > div {
                background-color: #162746;
                border: 1px solid #3b5f9f;
                border-radius: 12px;
                min-height: 44px;
                transition: border-color .18s ease, box-shadow .18s ease, transform .18s ease;
            }

            div[data-baseweb="input"] > div:focus-within,
            div[data-baseweb="textarea"] > div:focus-within,
            div[data-baseweb="select"] > div:focus-within {
                border-color: #e2b961;
                box-shadow: 0 0 0 1px #e2b961, 0 0 0 4px rgba(216, 171, 73, 0.18);
                transform: translateY(-1px);
            }

            input, textarea {
                color: #fcfcfc !important;
                font-size: 0.95rem !important;
            }

            textarea {
                line-height: 1.44 !important;
            }

            input::placeholder, textarea::placeholder {
                color: #a6badf !important;
            }

            div[data-testid="stMultiSelect"] span[data-baseweb="tag"] {
                background-color: rgba(62, 107, 193, 0.32);
                border: 1px solid rgba(216, 171, 73, 0.58);
                color: #e9f1ff;
            }

            div[data-testid="stButton"] > button,
            div[data-testid="stFormSubmitButton"] > button,
            div[data-testid="stDownloadButton"] > button {
                background:
                    radial-gradient(circle at 18% 15%, rgba(255, 239, 196, 0.32), rgba(255, 239, 196, 0) 44%),
                    linear-gradient(160deg, #2352ab 0%, #163d85 52%, #0f2b64 100%);
                color: #f8e5b1 !important;
                font-weight: 740;
                border: 1px solid rgba(235, 193, 101, 0.72);
                border-radius: 999px;
                padding: 0.58rem 1.36rem;
                box-shadow:
                    inset 0 1px 0 rgba(221, 236, 255, 0.38),
                    0 10px 24px rgba(0, 0, 0, 0.45),
                    0 0 0 1px rgba(28, 69, 145, 0.5);
                transition: transform .16s ease, filter .16s ease, box-shadow .16s ease, border-color .16s ease;
            }

            div[data-testid="stButton"] > button:hover,
            div[data-testid="stFormSubmitButton"] > button:hover,
            div[data-testid="stDownloadButton"] > button:hover {
                filter: brightness(1.09) saturate(1.12);
                transform: translateY(-2px) scale(1.006);
                border-color: rgba(244, 210, 134, 0.92);
                box-shadow:
                    inset 0 1px 0 rgba(229, 240, 255, 0.46),
                    0 14px 30px rgba(0, 0, 0, 0.52),
                    0 0 0 1px rgba(236, 194, 102, 0.36),
                    0 0 18px rgba(74, 123, 210, 0.35);
            }

            div[data-testid="stButton"] > button:focus-visible,
            div[data-testid="stFormSubmitButton"] > button:focus-visible,
            div[data-testid="stDownloadButton"] > button:focus-visible {
                outline: none;
                box-shadow:
                    inset 0 1px 0 rgba(229, 240, 255, 0.5),
                    0 0 0 2px rgba(242, 208, 128, 0.95),
                    0 0 0 6px rgba(76, 128, 217, 0.32),
                    0 12px 28px rgba(0, 0, 0, 0.5);
            }

            div[data-testid="stButton"] > button:disabled,
            div[data-testid="stFormSubmitButton"] > button:disabled,
            div[data-testid="stDownloadButton"] > button:disabled {
                opacity: 0.56;
                filter: grayscale(0.24) brightness(0.96);
                cursor: not-allowed;
                box-shadow: 0 6px 12px rgba(0, 0, 0, 0.28);
            }

            div[data-testid="stAlert"] {
                border-radius: 12px;
                border: 1px solid rgba(108, 157, 236, 0.42);
            }

            div[data-testid="stProgress"] > div > div {
                background-color: rgba(52, 78, 124, 0.44);
                border: 1px solid rgba(95, 142, 222, 0.34);
                border-radius: 999px;
                overflow: hidden;
            }

            div[data-testid="stProgress"] > div > div > div {
                background: linear-gradient(90deg, #e6bb66 0%, #5f93e8 48%, #2f64be 100%) !important;
                box-shadow: 0 0 12px rgba(216, 171, 73, 0.3);
                border-radius: 999px;
            }

            .preview-bloco {
                margin-top: 1.25rem;
                margin-bottom: 0.5rem;
                border-radius: 14px;
                border: 1px solid rgba(216, 171, 73, 0.4);
                background: linear-gradient(180deg, rgba(24, 43, 82, 0.9), rgba(11, 23, 46, 0.92));
                padding: 0.62rem 0.85rem;
                color: #f6dea6;
                font-weight: 650;
            }

            @media (max-width: 900px) {
                .main .block-container {
                    padding-top: 0.7rem;
                    padding-bottom: 1.4rem;
                }

                .hero-title {
                    font-size: 1.45rem;
                }

                .hero-subtitle {
                    font-size: 0.87rem;
                }

                .main div[data-testid="stVerticalBlockBorderWrapper"],
                div[data-testid="stForm"] {
                    padding: 0.95rem;
                    border-radius: 16px;
                }
            }
        </style>
        """,
        unsafe_allow_html=True,
    )


 # Renderiza o cabeçalho principal com informações de área, modelo e status da API.
def _render_cabecalho_moderno(area: str, modelo: str, api_configurada: bool) -> None:
    area_esc = html.escape(area or "Outro")
    modelo_esc = html.escape(modelo or "[PREENCHER]")
    status_label = "Chave API configurada" if api_configurada else "Chave API ausente"
    status_css = "status-ok" if api_configurada else "status-warn"

    st.markdown(
        f"""
        <section class="hero-shell">
            <h1 class="hero-title">Gerador de Petição Inicial</h1>
            <p class="hero-subtitle">
                Fluxo inteligente para montar petições iniciais com campos dinâmicos por área do direito.
            </p>
            <div class="hero-chip-row">
                <span class="hero-chip">Área: {area_esc}</span>
                <span class="hero-chip">Modelo: {modelo_esc}</span>
                <span class="hero-chip {status_css}">{status_label}</span>
            </div>
        </section>
        """,
        unsafe_allow_html=True,
    )


 # Renderiza um título de seção com índice visual e linha decorativa.
def _titulo_secao(texto: str) -> None:
    texto_limpo = (texto or "").strip()
    match = re.match(r"^(\d+)\)\s*(.*)$", texto_limpo)
    if match:
        indice = html.escape(match.group(1))
        titulo = html.escape(match.group(2).strip())
        indice_html = f'<span class="secao-indice">{indice}</span>'
    else:
        titulo = html.escape(texto_limpo)
        indice_html = ""

    st.markdown(
        f'<div class="secao-titulo">{indice_html}<span>{titulo}</span><span class="secao-linha"></span></div>',
        unsafe_allow_html=True,
    )


 # Renderiza o campo "Tipo da ação" com sugestões por área e opção de texto livre.
def _renderizar_campo_tipo_acao(area_direito: str) -> None:
    opcoes_area = TIPOS_ACAO_POR_AREA.get(area_direito, TIPOS_ACAO_POR_AREA["Outro"])
    kwargs_select: dict[str, Any] = {
        "label": "Tipo da ação *",
        "options": opcoes_area,
        "key": "tipo_acao",
        "placeholder": "Selecione ou digite para buscar/criar...",
        "accept_new_options": True,
        "help": "Escolha uma sugestão da área ou digite livremente.",
    }
    if "tipo_acao" not in st.session_state:
        kwargs_select["index"] = None
    st.selectbox(**kwargs_select)


 # Renderiza o campo "Rito/procedimento" com sugestões por área e opção de texto livre.
def _renderizar_campo_rito(area_direito: str) -> None:
    opcoes_rito = RITOS_POR_AREA.get(area_direito, RITOS_POR_AREA["Outro"])
    kwargs_select: dict[str, Any] = {
        "label": "Rito/Procedimento",
        "options": opcoes_rito,
        "key": "rito",
        "placeholder": "Selecione ou digite para buscar/criar...",
        "accept_new_options": True,
        "help": "Escolha uma sugestão da área ou digite livremente.",
    }
    if "rito" not in st.session_state:
        kwargs_select["index"] = None
    st.selectbox(**kwargs_select)


 # Retorna uma sugestão de foro/competência com base na área e rito escolhidos.
def _sugerir_foro_competente(area_direito: str, rito: str) -> str:
    area = (area_direito or "").strip()
    rito_norm = (rito or "").casefold()

    if area == "Previdenciário":
        if "juizado especial federal" in rito_norm:
            return "Sugestão de competência: Juizado Especial Federal (JEF), se cabível."
        return "Sugestão de competência: Justiça Federal (vara federal/previdenciária), salvo competência delegada."

    if area == "Direito da Saúde":
        return "Sugestão de competência: Vara Cível (plano/hospital privado) ou Vara da Fazenda Pública (SUS/ente público)."

    return ""


 # Aplica sugestões automáticas que reduzem erro de preenchimento sem impor campos.
def _aplicar_sugestoes_inteligentes(area_direito: str) -> None:
    if area_direito != "Direito da Saúde":
        return

    chave_urgencia = _chave_campo_area("Direito da Saude", "urgencia_laudo")
    urgencia_preenchida = bool(str(st.session_state.get(chave_urgencia, "")).strip())
    if not urgencia_preenchida:
        return

    # Autoativa uma vez quando há urgência informada; usuário pode desmarcar depois.
    if "tem_tutela_urgencia" not in st.session_state:
        st.session_state["tem_tutela_urgencia"] = True


load_dotenv()

st.set_page_config(page_title="Gerador de Peticao Inicial (Gemini)", layout="wide")
_aplicar_estilo_preto_dourado()
_restaurar_snapshot_formulario()
if st.session_state.get("area_direito") not in AREAS_DIREITO:
    st.session_state["area_direito"] = AREAS_DIREITO[0]

api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
gemini_model = os.getenv("GEMINI_MODEL", "gemini-2.5-flash").strip()
etapa_atual, etapa_idx = _menu_fluxo_lateral()
area_selecionada = st.session_state.get("area_direito", AREAS_DIREITO[0])
_aplicar_sugestoes_inteligentes(area_selecionada)
_render_cabecalho_moderno(area=area_selecionada, modelo=gemini_model, api_configurada=bool(api_key))
if not api_key:
    st.warning("Configure sua chave no .env (GEMINI_API_KEY ou GOOGLE_API_KEY) antes de gerar.")

voltar_etapa = False
avancar_etapa = False
gerar = False

with st.container(border=True):
    st.caption(f"Etapa atual: {etapa_atual}")
    modo_essencial = _modo_essencial_ativo()
    campos_obrigatorios_etapa = _campos_obrigatorios_da_etapa(etapa_atual)
    if campos_obrigatorios_etapa:
        rotulos = ", ".join(rotulo for _, rotulo in campos_obrigatorios_etapa)
        st.caption(f"Obrigatórios nesta etapa: {rotulos}")
    elif etapa_atual == "Pedidos":
        st.caption("Obrigatório nesta etapa: ao menos um pedido (base ou personalizado).")
    if modo_essencial:
        st.caption("Modo Essencial: campos opcionais avançados foram ocultados para reduzir o preenchimento.")

    if etapa_atual == "Contexto Processual":
        _titulo_secao("1) Contexto Processual")
        
        sugestao_foro = _sugerir_foro_competente(area_selecionada, str(st.session_state.get("rito", "")))
        foro_placeholder = "Ex.: Juizado Especial Federal" if area_selecionada == "Previdenciário" else "Ex.: Vara Cível / Fazenda Pública"

        ctx1, ctx2 = st.columns(2)
        with ctx1:
            _renderizar_campo_tipo_acao(area_selecionada)
        with ctx2:
            st.text_input("Comarca / UF *", key="comarca_uf", placeholder="Ex.: São Paulo/SP")

        ctx3, ctx4 = st.columns(2)
        with ctx3:
            _renderizar_campo_rito(area_selecionada)
        with ctx4:
            if not modo_essencial:
                st.text_input("Foro / Vara (opcional)", key="foro_vara", placeholder=foro_placeholder)

        if sugestao_foro:
            st.caption(sugestao_foro)

    elif etapa_atual == "Campos da Área":
        _titulo_secao("2) Campos Específicos da Área")
        _renderizar_bloco_area(area_selecionada)

    elif etapa_atual == "Partes":
        _titulo_secao("3) Partes")

        col1, col2 = st.columns(2)

        with col1:
            st.subheader("Autor")
            _renderizar_tipo_pessoa_parte("autor", area_selecionada)
            autor_tipo = str(st.session_state.get("autor_tipo_pessoa", "Pessoa Física"))
            autor_label_nome = "Nome *" if autor_tipo == "Pessoa Física" else "Razão social *"
            autor_label_doc = "CPF *" if autor_tipo == "Pessoa Física" else "CNPJ *"
            autor_placeholder_nome = "Ex.: Maria da Silva" if autor_tipo == "Pessoa Física" else "Ex.: Empresa XYZ LTDA"

            autor_id1, autor_id2 = st.columns(2)
            with autor_id1:
                st.text_input(autor_label_nome, key="autor_nome", placeholder=autor_placeholder_nome)
            with autor_id2:
                st.text_input(
                    autor_label_doc,
                    key="autor_doc",
                    placeholder="000.000.000-00 ou 00.000.000/0000-00",
                    max_chars=14 if autor_tipo == "Pessoa Física" else 18,
                    on_change=_aplicar_mascara_documento,
                    args=("autor_doc",),
                )
            if autor_tipo == "Pessoa Jurídica":
                st.button(
                    "Buscar dados da PJ por CNPJ",
                    key="btn_buscar_cnpj_autor",
                    on_click=_preencher_parte_com_cnpj,
                    args=("autor",),
                    use_container_width=True,
                )
                _exibir_feedback_cnpj("autor")
            if autor_tipo == "Pessoa Física":
                st.text_input(
                    "CEP (opcional)",
                    key="autor_cep",
                    placeholder="00000-000",
                    max_chars=9,
                    on_change=_aplicar_mascara_cep,
                    args=("autor_cep",),
                )
                st.button(
                    "Buscar endereço por CEP",
                    key="btn_buscar_cep_autor",
                    on_click=_preencher_endereco_por_cep,
                    args=("autor",),
                    use_container_width=True,
                )
                _exibir_feedback_cep("autor")
            else:
                st.caption("Para PJ, o CEP e o endereço podem ser preenchidos automaticamente pela consulta de CNPJ.")
            st.text_input(
                "Endereço *",
                key="autor_end",
                placeholder="Rua, número, bairro, cidade/UF",
            )
            if not modo_essencial:
                if autor_tipo == "Pessoa Física":
                    aut_pf1, aut_pf2, aut_pf3 = st.columns(3)
                    with aut_pf1:
                        st.text_input(
                            "Nacionalidade (opcional)",
                            key="autor_nacionalidade",
                            placeholder="Ex.: Brasileira",
                        )
                    with aut_pf2:
                        st.text_input(
                            "Estado civil (opcional)",
                            key="autor_estado_civil",
                            placeholder="Ex.: Solteira",
                        )
                    with aut_pf3:
                        st.text_input(
                            "Profissão (opcional)",
                            key="autor_profissao",
                            placeholder="Ex.: Professora",
                        )
                else:
                    aut_pj1, aut_pj2 = st.columns(2)
                    with aut_pj1:
                        st.text_input(
                            "Natureza jurídica (opcional)",
                            key="autor_natureza_juridica",
                            placeholder="Ex.: Pessoa jurídica de direito privado",
                        )
                    with aut_pj2:
                        st.text_input(
                            "Representante legal (opcional)",
                            key="autor_representante_legal",
                            placeholder="Ex.: João da Silva",
                        )
                st.text_area(
                    "Qualificação complementar do autor (opcional)",
                    key="autor_qualificacao",
                    height=90,
                    placeholder="Outras informações úteis de qualificação.",
                )

        with col2:
            st.subheader("Réu")
            _renderizar_tipo_pessoa_parte("reu", area_selecionada)
            reu_tipo = str(st.session_state.get("reu_tipo_pessoa", "Pessoa Jurídica"))
            reu_label_nome = "Nome *" if reu_tipo == "Pessoa Física" else "Razão social / Ente público *"
            reu_label_doc = "CPF *" if reu_tipo == "Pessoa Física" else "CNPJ *"
            reu_placeholder_nome = "Ex.: João da Silva" if reu_tipo == "Pessoa Física" else "Ex.: INSS / Município de Goiânia / Plano XYZ"

            reu_id1, reu_id2 = st.columns(2)
            with reu_id1:
                st.text_input(reu_label_nome, key="reu_nome", placeholder=reu_placeholder_nome)
            with reu_id2:
                st.text_input(
                    reu_label_doc,
                    key="reu_doc",
                    placeholder="000.000.000-00 ou 00.000.000/0000-00",
                    max_chars=14 if reu_tipo == "Pessoa Física" else 18,
                    on_change=_aplicar_mascara_documento,
                    args=("reu_doc",),
                )
            if reu_tipo == "Pessoa Jurídica":
                st.button(
                    "Buscar dados por CNPJ",
                    key="btn_buscar_cnpj_reu",
                    on_click=_preencher_parte_com_cnpj,
                    args=("reu",),
                    use_container_width=True,
                )
                _exibir_feedback_cnpj("reu")
            if reu_tipo == "Pessoa Física":
                st.text_input(
                    "CEP do réu (opcional)",
                    key="reu_cep",
                    placeholder="00000-000",
                    max_chars=9,
                    on_change=_aplicar_mascara_cep,
                    args=("reu_cep",),
                )
                st.button(
                    "Buscar endereço do réu por CEP",
                    key="btn_buscar_cep_reu",
                    on_click=_preencher_endereco_por_cep,
                    args=("reu",),
                    use_container_width=True,
                )
                _exibir_feedback_cep("reu")
            else:
                st.caption("Para PJ, o CEP e o endereço podem ser preenchidos automaticamente pela consulta de CNPJ.")
            st.text_input(
                "Endereço do réu *",
                key="reu_end",
                placeholder="Rua, número, bairro, cidade/UF",
            )
            if not modo_essencial:
                if reu_tipo == "Pessoa Física":
                    reu_pf1, reu_pf2, reu_pf3 = st.columns(3)
                    with reu_pf1:
                        st.text_input(
                            "Nacionalidade (opcional)",
                            key="reu_nacionalidade",
                            placeholder="Ex.: Brasileira",
                        )
                    with reu_pf2:
                        st.text_input(
                            "Estado civil (opcional)",
                            key="reu_estado_civil",
                            placeholder="Ex.: Casado",
                        )
                    with reu_pf3:
                        st.text_input(
                            "Profissão (opcional)",
                            key="reu_profissao",
                            placeholder="Ex.: Comerciante",
                        )
                else:
                    reu_pj1, reu_pj2 = st.columns(2)
                    with reu_pj1:
                        st.text_input(
                            "Natureza jurídica (opcional)",
                            key="reu_natureza_juridica",
                            placeholder="Ex.: Autarquia federal / Pessoa jurídica de direito privado / Pessoa jurídica de direito público",
                        )
                    with reu_pj2:
                        st.text_input(
                            "Representante legal (opcional)",
                            key="reu_representante_legal",
                            placeholder="Ex.: Procuradoria Federal Especializada / Prefeito Municipal / Diretor",
                        )
                st.text_area(
                    "Qualificação complementar do réu (opcional)",
                    key="reu_qualificacao",
                    height=90,
                    placeholder="Outras informações úteis de qualificação.",
                )

        if not modo_essencial:
            st.text_area(
                "Partes adicionais (opcional, uma por linha)",
                key="partes_adicionais_raw",
                height=90,
                placeholder="Ex.: Litisconsorte ativo | Nome | CPF/CNPJ | Endereço",
            )

    elif etapa_atual == "Fatos e Provas":
        _titulo_secao("4) Narrativa Fática")
        st.text_area(
            "Fatos principais *",
            key="fatos",
            height=180,
            placeholder="Descreva os fatos de forma cronológica e objetiva.",
        )
        if not modo_essencial:
            st.text_area(
                "Cronologia detalhada (opcional, um evento por linha)",
                key="cronologia_raw",
                height=100,
                placeholder="Ex.: 10/01/2026 - Contrato assinado",
            )
            provas_sugeridas_area = PROVAS_SUGERIDAS_POR_AREA.get(area_selecionada, PROVAS_SUGERIDAS_POR_AREA["Outro"])
            st.multiselect(
                "Provas sugeridas para esta área",
                provas_sugeridas_area,
                key="provas_sugeridas",
                help="Selecione as provas já disponíveis no caso.",
            )
            st.text_area(
                "Documentos e provas adicionais (opcional, um item por linha)",
                key="provas_raw",
                height=100,
                placeholder="Ex.: Contrato, comprovantes de pagamento, trocas de e-mail",
            )

    elif etapa_atual == "Fundamentação":
        _titulo_secao("5) Fundamentação Jurídica")
        st.text_area(
            "Teses jurídicas (linha de argumentação) *",
            key="teses_juridicas",
            height=130,
            placeholder="Quais pontos jurídicos devem ser defendidos na petição.",
        )

        if not modo_essencial:
            fun1, fun2 = st.columns(2)
            with fun1:
                st.multiselect("Temas jurídicos comuns", TEMAS_JURIDICOS_COMUNS, key="temas_comuns")
            with fun2:
                st.text_area(
                    "Fundamentos legais (opcional, um por linha)",
                    key="fundamentos_legais_raw",
                    height=120,
                    placeholder="Ex.: Art. 186 do CC; Art. 6, VIII, do CDC",
                )

            st.text_area(
                "Temas jurídicos adicionais (opcional, um por linha)",
                key="temas_custom_raw",
                height=90,
                placeholder="Ex.: Teoria do adimplemento substancial",
            )

    elif etapa_atual == "Pedidos":
        _titulo_secao("6) Pedidos")
        pedidos_base_exibir = _pedidos_base_exibicao()
        if modo_essencial:
            st.multiselect("Pedidos base", pedidos_base_exibir, key="pedidos_base")
        else:
            ped1, ped2 = st.columns(2)
            with ped1:
                st.multiselect("Pedidos base", pedidos_base_exibir, key="pedidos_base")
            with ped2:
                st.text_area(
                    "Pedidos personalizados (um por linha)",
                    key="pedidos_custom_raw",
                    height=150,
                    placeholder="Ex.: Condenação ao pagamento de R$ 15.000,00 a título de dano moral",
                )
        st.caption("Tutela de urgência e justiça gratuita são definidas na etapa final para evitar duplicidade.")

    elif etapa_atual == "Finalização e Geração":
        _titulo_secao("7) Estrutura da Peça e Parâmetros Finais")
        fim1, fim2 = st.columns(2)

        with fim1:
            st.multiselect(
                "Seções sugeridas para organizar a petição",
                SECOES_SUGERIDAS,
                default=[
                    "Dos fatos",
                    "Do direito",
                    "Dos pedidos",
                    "Do valor da causa",
                    "Dos requerimentos finais",
                ],
                key="secoes_sugeridas",
            )
            if not modo_essencial:
                st.text_area(
                    "Seções extras personalizadas (opcional, uma por linha)",
                    key="secoes_extras_raw",
                    height=100,
                    placeholder="Ex.: Da inversão do ônus da prova",
                )

        with fim2:
            st.text_input(
                "Valor da causa (se souber) *",
                key="valor_causa",
                placeholder="Digite apenas números. Ex.: 125000 -> R$ 1.250,00",
                on_change=_aplicar_mascara_moeda,
                args=("valor_causa",),
            )
            st.radio(
                "Nível de detalhamento",
                NIVEIS_DETALHAMENTO,
                index=1,
                horizontal=True,
                key="nivel_detalhamento",
            )
            st.checkbox("Incluir pedido de tutela de urgência", key="tem_tutela_urgencia")
            st.checkbox("Incluir pedido de justiça gratuita", key="tem_gratuidade")
            st.checkbox("Incluir prioridade de tramitação", key="tem_prioridade")
            st.checkbox("Manifestar interesse em audiência de conciliação", key="quer_audiencia", value=True)
            if area_selecionada == "Direito da Saúde":
                chave_urgencia = _chave_campo_area("Direito da Saude", "urgencia_laudo")
                urgencia_laudo = str(st.session_state.get(chave_urgencia, "")).strip()
                if urgencia_laudo and not st.session_state.get("tem_tutela_urgencia", False):
                    st.info("Há urgência médica informada. Sugestão: incluir pedido de tutela de urgência.")
                    if st.button("Aplicar sugestão de tutela", key="btn_aplicar_sugestao_tutela"):
                        st.session_state["tem_tutela_urgencia"] = True
                        st.rerun()

        if not modo_essencial:
            st.text_area(
                "Observações estratégicas para a redação (opcional)",
                key="obs_estrategicas",
                height=100,
                placeholder="Diretrizes de linguagem, foco, riscos e pontos sensíveis.",
            )

            st.markdown("#### Modelo de Referência (opcional)")
            st.caption("Envie um modelo para orientar estilo e estrutura. Formatos: .txt, .md, .docx")
            arquivo_modelo_referencia = st.file_uploader(
                "Arquivo do modelo",
                type=["txt", "md", "docx"],
                key="modelo_referencia_upload",
                help="O sistema usa o arquivo apenas como referencia de redacao, sem substituir os dados do caso.",
            )
            _processar_modelo_referencia(arquivo_modelo_referencia)

            erro_modelo = str(st.session_state.get("_modelo_referencia_erro", "")).strip()
            if erro_modelo:
                st.error(erro_modelo)
            else:
                nome_modelo = str(st.session_state.get("modelo_referencia_nome", "")).strip()
                texto_modelo = str(st.session_state.get("modelo_referencia_texto", "")).strip()
                truncado_modelo = bool(st.session_state.get("modelo_referencia_truncado", False))
                if nome_modelo and texto_modelo:
                    sufixo = " (trecho truncado para caber no prompt)" if truncado_modelo else ""
                    st.success(f"Modelo carregado: {nome_modelo}{sufixo}")
                    st.text_area(
                        "Prévia extraída do modelo",
                        value=texto_modelo,
                        height=140,
                        disabled=True,
                    )

        st.markdown("#### Dados do Advogado (para o fechamento da peça)")
        adv1, adv2, adv3 = st.columns([2, 1, 1])
        with adv1:
            st.text_input(
                "Nome do advogado(a)",
                key="advogado_nome",
                placeholder="Ex.: João da Silva",
            )
        with adv2:
            st.text_input(
                "OAB/UF",
                key="advogado_oab_uf",
                placeholder="Ex.: SP",
                max_chars=2,
            )
        with adv3:
            st.text_input(
                "Número da OAB",
                key="advogado_oab_num",
                placeholder="Ex.: 123456",
            )

    if etapa_atual == "Finalização e Geração":
        nav1, nav2 = st.columns(2)
        with nav1:
            voltar_etapa = st.button("Voltar", key="btn_voltar_final")
        with nav2:
            gerar = st.button("Gerar petição", key="btn_gerar")
    else:
        nav1, nav2 = st.columns(2)
        with nav1:
            voltar_etapa = st.button("Voltar", key="btn_voltar", disabled=etapa_idx == 0)
        with nav2:
            avancar_etapa = st.button("Avançar", key="btn_avancar")

_salvar_snapshot_formulario()

if "peticao_texto" not in st.session_state:
    st.session_state.peticao_texto = ""

if voltar_etapa:
    _definir_etapa_idx(etapa_idx - 1)
    st.rerun()

if avancar_etapa:
    faltantes_etapa = _validar_etapa(etapa_atual)
    if faltantes_etapa:
        itens = "\n".join(f"- {item}" for item in faltantes_etapa)
        st.error(f"Preencha os campos obrigatórios antes de avançar:\n{itens}")
    else:
        _definir_etapa_idx(etapa_idx + 1)
        st.rerun()

if gerar:
    _restaurar_snapshot_formulario()
    faltantes_geracao = _validar_essenciais_para_geracao()
    if faltantes_geracao:
        itens = "\n".join(f"- {item}" for item in faltantes_geracao)
        st.error(f"Não é possível gerar ainda. Campos principais obrigatórios pendentes:\n{itens}")
    else:
        dados = _coletar_payload()
        prompt = montar_prompt(dados)

        with st.spinner("Gerando a peticao..."):
            try:
                texto = gerar_peticao(prompt, model=gemini_model)
                st.session_state.peticao_texto = texto
            except GeminiServiceError as exc:
                st.error(str(exc))
            except Exception as exc:  # pragma: no cover
                st.error(f"Erro inesperado ao gerar peticao: {exc}")

if st.session_state.peticao_texto:
    st.markdown('<div class="preview-bloco">Prévia da petição gerada</div>', unsafe_allow_html=True)
    st.text_area("Texto gerado", st.session_state.peticao_texto, height=420)
    nome_arquivo_docx = _nome_arquivo_docx(st.session_state.get("autor_nome", ""))
    nome_arquivo_pdf = _nome_arquivo_pdf(st.session_state.get("autor_nome", ""))

    docx_bytes = texto_para_docx_bytes(
        titulo="PETICAO INICIAL",
        texto=st.session_state.peticao_texto,
    )
    pdf_bytes = texto_para_pdf_bytes(
        titulo="PETICAO INICIAL",
        texto=st.session_state.peticao_texto,
    )

    down1, down2 = st.columns(2)
    with down1:
        st.download_button(
            "Baixar .docx",
            data=docx_bytes,
            file_name=nome_arquivo_docx,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    with down2:
        st.download_button(
            "Baixar .pdf",
            data=pdf_bytes,
            file_name=nome_arquivo_pdf,
            mime="application/pdf",
        )
